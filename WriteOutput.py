def WriteOutput(files_dict):
    # Import Modules
    import glob
    import os
    import docx
    from docx.shared import Pt 
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_BREAK
    from docx.enum.section import WD_ORIENT, WD_SECTION
    import matplotlib.pyplot as plt
    from docx.shared import Inches

    # Set Paths
    home = os.getcwd()
    data_path = r'C:\Users\bhearley\Box\Lab Infrastructure Data Collection\Final'

    # Utility Function
    def change_orientation():
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        return new_section

    # Create the Document
    doc = docx.Document() 

    # Create the Title Page
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run('NASA GRC Lab Infrastructure Data')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(18)
    run1.bold = True

    # Loop Through Divisions
    divisions = list(files_dict.keys())
    divisions.sort()

    for d in range(len(divisions)):
        # Start on New Page
        doc.add_page_break()

        run_lab1 = doc.add_paragraph().add_run(divisions[d])
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(14)
        run_lab1.bold = True

        # Get list of branches
        branches = list(files_dict[divisions[d]].keys())
        branches.sort()

        for b in range(len(branches)):
            run_lab1 = doc.add_paragraph().add_run(branches[b])
            run_lab1.font.name = 'Times New Roman'
            run_lab1.font.size = Pt(12)
            run_lab1.bold = True

            # Get List of files
            files = files_dict[divisions[d]][branches[b]]
            files.sort()

            for q in range(len(files)):
                # Read the Text File
                with open(os.path.join(data_path,files[q])) as f:
                    lines = f.readlines()

                
                # -- Laboratory/Capability Name
                key = 'Laboratory/Capability Name:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(18)
                run_lab1.bold = True

                # HEADER: Laboratory/Capability Information
                run_lab1 = doc.add_paragraph().add_run('Laboratory/Capability Information')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Point of Contact
                key = 'Point of Contact:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Point of Contact
                key = 'Branch:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Description
                key = 'Laboratory/Capability Description:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Website
                key = 'Laboratory/Capability Website:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Challenges in sustaining this laboratory/capability
                key = 'Challenges in sustaining this laboratory/capability:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Age (yrs):
                key = 'Age (yrs):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Condition:
                key = 'Condition:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

    
                # -- Asset Table
                key = 'Number of Assets:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_assets  = int(val)
        
                data = ''
                for k in range(line_num+2,line_num+2+num_assets):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_assets):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_assets > 0:
                    change_orientation()

                    run_lab1 = doc.add_paragraph().add_run('Assets:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    table = doc.add_table(rows=1, cols=11) 
                    row = table.rows[0].cells 
                    row[0].text = 'Asset Name'
                    row[1].text = 'Location (Bldg/Rm)'
                    row[2].text = 'Age (yrs)'
                    row[3].text = 'Asset Date of Entry'
                    row[4].text = 'Expected Date of Obsolescence'
                    row[5].text = 'Asset Condition'
                    row[6].text = 'Replacement Cost ($)'
                    row[7].text = 'Impact to Capability if Lost'
                    row[8].text = 'Associated Software/Required OS'
                    row[9].text = 'IT Hardware Repalcement?'
                    row[10].text = 'Part or Full Replacement?'

                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(11):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'
                    run_lab1 = doc.add_paragraph().add_run('')

                    change_orientation()

                # -- Sustainment Funding Source:
                key = 'Sustainment Funding Source:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Funding Table
                key = 'Number of Funding Sources:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_fund  = int(val)

                data = ''
                for k in range(line_num+2,line_num+2+num_fund):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_fund):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)
    
                if num_fund > 0:
                    run_lab1 = doc.add_paragraph().add_run('Funding Sources:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=4) 
                    row = table.rows[0].cells 
                    row[0].text = 'Funding Source'
                    row[1].text = 'Funding Start Date'
                    row[2].text = 'Funding End Date'
                    row[3].text = 'Funding Amount per Year ($)'


                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(4):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Current Mission/Project Utilization
                run_lab1 = doc.add_paragraph().add_run('Current Mission/Project Utilization')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # Project Table
                key = 'Number of Projects:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_proj  = int(val)

                data = ''
                for k in range(line_num+2,line_num+2+num_proj):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_proj):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_proj > 0:
                    run_lab1 = doc.add_paragraph().add_run('Projects:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=5) 
                    row = table.rows[0].cells 
                    row[0].text = 'Mission/Project Name'
                    row[1].text = 'WBS Number'
                    row[2].text = 'Project Use (%)'
                    row[3].text = 'Risk to Project'
                    row[4].text = 'Impact if Laboratory/Capability is Lost'


                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(5):
                            if k == 1:
                                row[k].text = data_all[j][k][0:6]
                            else:
                                row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(len(data_all)):
                        labels.append(data_all[j][0])
                        vals.append(float(data_all[j][2]))

                    run_lab1 = doc.add_paragraph().add_run('')
                    fig, ax = plt.subplots()
                    ax.pie(vals, labels=labels, autopct='%1.0f%%')
                    plt.savefig(os.path.join(data_path,'Project_chart_' + str(q)+'.png'))
                    doc.add_picture(os.path.join(data_path,'Project_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # HEADER: Utilization History/Impact
                run_lab1 = doc.add_paragraph().add_run('Utilization History/Impact')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- History of capability utilization
                key = 'History of capability utilization:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Major impact and contributions this capability has made possible:
                key = 'Major impact and contributions this capability has made possible:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # HEADER: History of Down Time Due to Maintenance or Failure
                run_lab1 = doc.add_paragraph().add_run('History of Down Time Due to Maintenance or Failure')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Read Down Time Table
                key = 'Number of Failures:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_dt  = int(val)
        
                data = ''
                for k in range(line_num+2,line_num+2+num_dt):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_dt):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_dt > 0:
                    run_lab1 = doc.add_paragraph().add_run('Previous Laboratory/Asset Failures:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=5) 
                    row = table.rows[0].cells 
                    row[0].text = 'Asset'
                    row[1].text = 'Start Date'
                    row[2].text = 'Time Down'
                    row[3].text = 'Time Down Unit'
                    row[4].text = 'Additional Notes'

                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(5):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Cost
                run_lab1 = doc.add_paragraph().add_run('Cost')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Estimated Cost to Replace Entire Laboratory/Capability ($):
                key = 'Estimated Cost to Replace Entire Laboratory/Capability ($):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
    
        
                # -- Cost of Service Contracts ($):
                key = 'Cost of Service Contracts ($):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Annual Cost to Operate and Sustain the Lab ($/yr):
                key = 'Annual Cost to Operate and Sustain the Lab ($/yr):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Cost of Service Contracts ($):
                key = 'Incurred Cost For Downtime ($/yr):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Read Divisons Table
                key = 'Number of Divisions (Labor Costs):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_div  = int(val)

        
                data = ''
                for k in range(line_num+2,line_num+2+num_div):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_div):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_div > 0:
                    run_lab1 = doc.add_paragraph().add_run('Directorate Labor Division:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=2) 
                    row = table.rows[0].cells 
                    row[0].text = 'Directorate'
                    row[1].text = 'Labor Division (%)'
                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(2):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'
                    run_lab1 = doc.add_paragraph().add_run('')

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(len(data_all)):
                        labels.append(data_all[j][0])
                        vals.append(float(data_all[j][1]))

                    run_lab1 = doc.add_paragraph().add_run('')
                    fig, ax = plt.subplots()
                    ax.pie(vals, labels=labels, autopct='%1.0f%%')
                    plt.savefig(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'))
                    doc.add_picture(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # Start on New Page
                doc.add_page_break()




    # Save the Document
    doc.save(os.path.join(data_path,'Lab Data Output - Filtered.docx')) 


    temp=1